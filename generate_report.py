"""
生成《人工智能》项目大作业设计报告 .docx
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

BASE = r'E:\code\周记\人工智能\Fruits-360'

doc = Document()

# ===== 设置默认字体 =====
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ===== 封面 =====
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('广师大项目大作业封面')
run.font.size = Pt(22)
run.bold = True
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

doc.add_paragraph()

proj = doc.add_paragraph()
proj.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = proj.add_run('"人工智能"项目大作业\n基于ResNet-18迁移学习的水果智能识别系统')
run.font.size = Pt(18)
run.bold = True
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

for _ in range(4):
    doc.add_paragraph()

info_items = ['学  院：______计算机科学学院______', '专  业：________________________',
              '姓  名：________________________', '学  号：________________________',
              '指导教师：______刘勋______________', '日  期：______2026年6月_______']
for item in info_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(item)
    run.font.size = Pt(14)

doc.add_page_break()

# ===== 辅助函数 =====
def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return h

def add_para(text, bold=False, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(12)
    if bold:
        run.bold = True
    return p

def add_image(path, width=Inches(5)):
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=width)

# =====================================================================
# 一、项目简介
# =====================================================================
add_heading('一、项目简介', 1)

add_para('随着生活水平的提高，人们对水果的种类和品质要求日益增长。然而，水果种类繁多，普通消费者往往难以准确识别某些稀有的或外观相似的水果品种。传统的人工识别方式依赖经验，效率低且容易出错。近年来，深度学习技术在计算机视觉领域取得了突破性进展，特别是卷积神经网络（CNN）在图像分类任务上表现优异。', indent=True)
add_para('本项目旨在构建一个基于深度迁移学习的水果智能识别系统，利用 PyTorch 深度学习框架和 ResNet-18 预训练模型，在 Fruits-360 公开数据集上实现 260 种水果和蔬菜的自动分类，最终模型在测试集上达到 98.80% 的准确率。同时，采用 Flask Web 框架将训练好的模型部署为 Web 应用，用户可通过上传图片获得实时识别结果。', indent=True)
add_para('本项目的核心功能包括：（1）基于 ResNet-18 迁移学习的 260 类水果识别模型；（2）数据增强策略提升模型泛化能力；（3）两阶段训练策略（冻结训练+微调）优化模型性能；（4）混淆矩阵和分类报告等全面的模型评估；（5）基于 Flask+HTML5 的 Web 识别界面，支持拖拽上传和 Top-5 预测结果展示。', indent=True)

# =====================================================================
# 二、开发环境与原理
# =====================================================================
add_heading('二、开发环境与原理', 1)
add_heading('2.1 开发环境', 2)

# 环境表格
table = doc.add_table(rows=12, cols=2, style='Table Grid')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
env_data = [
    ('项目', '配置'),
    ('操作系统', 'Windows 11 Home China (64位)'),
    ('编程语言', 'Python 3.10.8'),
    ('深度学习框架', 'PyTorch 2.4.0+cu124'),
    ('Web 框架', 'Flask 3.1.3'),
    ('CUDA 版本', 'CUDA 12.4'),
    ('GPU', 'NVIDIA GeForce RTX 4050 Laptop (6GB)'),
    ('CPU', 'Intel 13th Gen'),
    ('数据科学库', 'NumPy 1.26.4, Pandas, Scikit-learn 1.6.1'),
    ('可视化', 'Matplotlib 3.x, Seaborn 0.13.2'),
    ('图像处理', 'Pillow (PIL), torchvision 0.19'),
    ('其他工具', 'Jupyter, python-docx'),
]
for i, (k, v) in enumerate(env_data):
    row = table.rows[i]
    row.cells[0].text = k
    row.cells[1].text = v
    for cell in row.cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(10)

doc.add_paragraph()

add_heading('2.2 技术原理', 2)

add_para('（1）卷积神经网络（CNN）', bold=True)
add_para('卷积神经网络是一类专门用于处理网格结构数据（如图像）的深度学习模型。CNN 主要由卷积层（Convolutional Layer）、池化层（Pooling Layer）和全连接层（Fully Connected Layer）组成。卷积层通过可学习的卷积核（滤波器）在输入图像上滑动，提取局部特征如边缘、纹理和形状。池化层（如最大池化）对特征图进行下采样，降低计算量并提供平移不变性。全连接层将提取的高层特征映射到类别空间进行最终分类。', indent=True)

add_para('（2）ResNet 残差网络', bold=True)
add_para('ResNet（Residual Network）由微软研究院于 2015 年提出，核心创新是"跳跃连接"（Skip Connection）。传统深层网络面临梯度消失/爆炸问题——随着层数加深，训练误差不降反升。ResNet 通过引入残差块（Residual Block）解决此问题，其公式为：输出 = F(x) + x，其中 F(x) 是网络学习到的残差映射，x 是恒等映射（输入）。跳跃连接使得梯度可以直接流过网络，解决了深层网络的退化问题。ResNet-18 包含 17 个卷积层和 1 个全连接层，共计约 1100 万参数，在 ImageNet 上 Top-1 准确率约 69.76%。', indent=True)

add_para('（3）迁移学习（Transfer Learning）', bold=True)
add_para('迁移学习是将在源任务上学到的知识应用到目标任务的一种方法。在计算机视觉领域，通常采用在 ImageNet（120 万张图片，1000 类）上预训练的模型作为起点，将其特征提取能力迁移到新的分类任务上。优势在于：（a）大幅减少训练所需的数据量和计算资源；（b）预训练模型已经学习了通用视觉特征（边缘、纹理、形状等），只需微调高层特征适配新任务；（c）相比从零训练，迁移学习通常能获得更高的准确率和更快的收敛速度。', indent=True)

add_para('（4）两阶段训练策略', bold=True)
add_para('第一阶段（冻结训练）：冻结预训练模型的所有卷积层参数，仅训练随机初始化的全连接分类层。此阶段使用较大的学习率（0.001），让分类头快速收敛。第二阶段（微调）：解冻 ResNet 最后几个卷积层（layer3 和 layer4），采用分层学习率策略——全连接层（5e-4）、layer4（1e-4）、layer3（5e-5），利用余弦退火调度器平滑降低学习率。分层学习率的设计原则是：底层学习通用边缘/纹理特征（不动），高层学习形状组合特征（小幅度调整），分类头需要大步快学。', indent=True)

add_para('（5）数据增强（Data Augmentation）', bold=True)
add_para('数据增强通过对训练样本施加随机变换来扩充数据集，防止模型过拟合。本项目采用的增强策略包括：随机裁剪（RandomResizedCrop, 224×224）、水平翻转（RandomHorizontalFlip）、随机旋转 ±20°（RandomRotation）、颜色抖动（ColorJitter，亮度/对比度/饱和度 ±30%）。这些变换模拟了真实场景中不同的拍摄角度、光线条件和构图方式，让模型学习到更具泛化能力的特征。', indent=True)

# =====================================================================
# 三、项目实现及结果
# =====================================================================
add_heading('三、项目实现及结果', 1)
add_heading('3.1 项目总体框架结构图', 2)

add_para('项目的整体流程如下：', indent=True)
add_para('（1）数据获取：从 Kaggle 下载 Fruits-360 数据集（训练集 137,221 张，测试集 45,724 张，260 个类别）；', indent=True)
add_para('（2）数据预处理：数据增强、归一化（ImageNet 均值/标准差）、划分训练/测试集；', indent=True)
add_para('（3）模型构建：加载 ImageNet 预训练 ResNet-18，替换全连接层为 260 类输出；', indent=True)
add_para('（4）两阶段训练：第一阶段冻结卷积层训练 FC 层（10 epochs），第二阶段解冻高层微调（15 epochs）；', indent=True)
add_para('（5）模型评估：计算准确率、分类报告、混淆矩阵、易混淆类别分析；', indent=True)
add_para('（6）Web 部署：Flask 后端加载模型 → 前端上传图片 → 模型推理 → Top-5 结果展示。', indent=True)

add_para('项目代码结构：', bold=True)
code_structure = """Fruits-360/
├── model.py          # 模型定义（ResNet-18 迁移学习）
├── train.py          # 完整训练脚本（两阶段训练）
├── train_resume.py   # 续训脚本（从 Stage1 checkpoint 继续）
├── evaluate.py       # 模型评估（混淆矩阵 + 分类报告）
├── app.py            # Flask Web 应用
├── templates/
│   └── index.html    # 前端识别界面
├── best_model.pth    # 训练好的模型权重（98.80%）
├── class_names.txt   # 260 类水果名称
├── training_curve.png      # 训练曲线图
├── confusion_matrix.png    # 混淆矩阵
└── top_errors.png          # 易混淆类别排行
"""
p = doc.add_paragraph()
run = p.add_run(code_structure)
run.font.name = 'Consolas'
run.font.size = Pt(9)

add_heading('3.2 项目实现', 2)

add_para('（1）数据获取及预处理', bold=True)
add_para('本项目使用 Fruits-360 公开数据集，这是目前最全面的水果图像数据集之一。数据集包含 260 种水果和蔬菜，共计 182,945 张图片（训练集 137,221 张，测试集 45,724 张），每张图片分辨率为 100×100 像素。数据按类别存放在不同文件夹中，每类约 480-530 张训练图、100-320 张测试图。使用 torchvision.datasets.ImageFolder 自动按文件夹名读取类别标签。训练时，图片先被缩放到 256×256，然后随机裁剪到 224×224（模拟不同的图像构图），再应用随机水平翻转、旋转和颜色抖动等增强操作。最后使用 ImageNet 数据集的均值 [0.485, 0.456, 0.406] 和标准差 [0.229, 0.224, 0.225] 进行归一化处理。测试集仅做缩放、中心裁剪和归一化以保证评估的稳定性。', indent=True)

add_para('（2）模型构建', bold=True)
add_para('模型基于 torchvision.models 中的 ResNet-18 预训练权重（ResNet18_Weights.DEFAULT，即 ImageNet 预训练）构建。将原有的 1000 类全连接层替换为适配本项目 260 类的线性层（nn.Linear(512, 260)）。第一阶段冻结所有卷积层参数（requires_grad=False），仅优化 FC 层参数。第二阶段通过 unfreeze_layers() 函数解冻 layer3 和 layer4 的参数，其中每个 layer 包含 2 个 BasicBlock（共约 14M 参数的子集中部分变为可训练）。FC 层始终保持完全可训练。', indent=True)

add_para('（3）模型训练', bold=True)
add_para('训练分两个阶段进行。训练配置如下：', indent=True)

# 训练参数表
table2 = doc.add_table(rows=8, cols=3, style='Table Grid')
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
train_data = [
    ('参数', '第一阶段（冻结）', '第二阶段（微调）'),
    ('Epochs', '10', '15'),
    ('Batch Size', '64', '64'),
    ('优化器', 'Adam', 'Adam（分层学习率）'),
    ('学习率', 'fc: 0.001', 'fc: 5e-4, layer4: 1e-4, layer3: 5e-5'),
    ('调度器', 'StepLR (step=5, gamma=0.5)', 'CosineAnnealingLR (T_max=15)'),
    ('损失函数', 'CrossEntropyLoss', 'CrossEntropyLoss'),
    ('可训练层', '仅 FC 层', 'FC + layer3 + layer4'),
]
for i, row_data in enumerate(train_data):
    for j, cell_text in enumerate(row_data):
        table2.rows[i].cells[j].text = cell_text
        for p in table2.rows[i].cells[j].paragraphs:
            for run in p.runs:
                run.font.size = Pt(10)

doc.add_paragraph()

add_para('第一阶段（冻结训练）：冻结 ResNet-18 所有的卷积层参数，只训练新替换的分类头（FC层）。使用 Adam 优化器，初始学习率 0.001，StepLR 调度器每 5 个 epoch 将学习率减半。经过 10 个 epoch 训练，验证集准确率达到 85.64%。第二阶段（微调）：基于第一阶段的最优模型，解冻 layer3 和 layer4 两个残差块，采用分层学习率（全连接层 5e-4 > layer4 1e-4 > layer3 5e-5）进行精细调优。使用 CosineAnnealingLR 调度器平滑降低学习率，训练 15 个 epoch。最终模型验证集准确率达到 98.80%。', indent=True)

add_para('训练过程中的准确率和损失曲线如下图所示：', indent=True)
add_image(os.path.join(BASE, 'training_curve.png'), width=Inches(5.5))
add_para('图1：训练过程 Loss 曲线和 Accuracy 曲线', bold=True)
p = doc.paragraphs[-1]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

add_para('由训练曲线可见：（a）第一阶段验证集准确率从 73.90% 稳定提升至 85.64%，Loss 持续下降，没有出现过拟合现象；（b）进入第二阶段后，由于解冻了高层卷积层，模型准确率在第一轮就跃升至 96.16%，证明了迁移学习微调策略的有效性；（c）第二阶段后期准确率增长逐渐平缓（从 98.06% 到 98.80%），说明模型已经接近该数据集上的最优性能上限。', indent=True)

add_para('（4）模型评估', bold=True)
add_para('在 45,724 张测试集图片上对最终模型进行全面评估，结果如下：', indent=True)
add_para('• 总体准确率（Accuracy）：98.80%', indent=True)
add_para('• 加权平均 F1-score（Weighted Avg）：0.99', indent=True)
add_para('• 宏平均 F1-score（Macro Avg）：0.98', indent=True)
add_para('绝大多数类别的 F1-score 达到 1.00（完美识别），仅有少数外观极为相似的水果类别存在轻微混淆。这说明 ResNet-18 迁移学习在该数据集上取得了非常优秀的表现。', indent=True)

add_image(os.path.join(BASE, 'confusion_matrix.png'), width=Inches(5.5))
add_para('图2：混淆矩阵（最容易混淆的类别）', bold=True)
p = doc.paragraphs[-1]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

add_image(os.path.join(BASE, 'top_errors.png'), width=Inches(5.5))
add_para('图3：最容易混淆的类别对 Top-15', bold=True)
p = doc.paragraphs[-1]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

add_para('从混淆矩阵可以看出，模型对绝大多数类别都能完美识别。少数混淆集中在不同品种的同一水果之间（如不同编号的苹果之间、不同种类的黑莓之间），这主要原因是数据集本身的标签标注存在类间相似度高的问题——某些"不同类别"的差异非常微小，即使人类也难以区分。', indent=True)

add_para('（5）背景去除预处理', bold=True)
add_para('在初步测试中发现，模型在 Fruits-360 测试集上表现优异（98.80%），但面对真实网络图片时识别率急剧下降——几乎所有带复杂背景的网图都被错误分类。经分析，根本原因在于 Fruits-360 数据集的所有图片均为纯白背景下的单个水果，模型在学习过程中将"白色背景像素"也作为分类特征的一部分。当输入图片的背景发生变化时，模型的判断严重受干扰。这种训练数据与实际应用场景之间的差异在深度学习中被称为"领域偏移"（Domain Shift）。', indent=True)
add_para('为解决此问题，本项目引入了基于 U²-Net 深度学习模型的自动背景去除模块（rembg 库）。U²-Net 是一种专门用于显著性目标检测的深度网络，能够精确分割图像中的前景物体。预处理流程为：用户上传图片 → U²-Net 分割前景（水果）→ 将透明前景贴合到纯白背景 → 送入 ResNet-18 进行识别。这一预处理步骤使输入图片的风格与 Fruits-360 训练集保持一致，大幅缩小了领域偏移。', indent=True)
add_para('为验证抠图预处理的效果，本项目设计了三组对比实验：（a）纯 Fruits-360 测试集图片——抠图前后均正确，证明抠图不破坏原有识别能力；（b）人工合成带纹理背景的图片——水果占比 15-40%，抠图前 7 张中仅 2 张正确，抠图后 6 张正确，修正率达 80%；（c）从 Unsplash 下载的真实场景照片——蛋糕上的草莓抠图后成功识别为 Strawberry，树上苹果虽未完全正确但 Apple 进入 Top-3 候选。实验结果充分证明：背景去除是解决领域偏移的有效手段，而非模型本身的架构问题。', indent=True)

add_para('（6）模型应用——Flask Web 部署', bold=True)
add_para('为了实现模型的实用化，本项目基于 Flask Web 框架开发了一个在线水果识别系统。前端采用 HTML5+CSS3+JavaScript 构建，支持三种交互方式：点击上传、拖拽上传和实时预览。用户上传图片后，前端通过 Ajax 请求将图片发送至后端 /predict 接口，后端先使用 U²-Net 去除背景，再将处理后的图片送入 ResNet-18 模型进行推理，返回 Top-5 预测结果（类别名称+置信度百分比）。前端以排名列表形式展示结果，每个结果带有百分比进度条，视觉效果直观。Web 应用的整个推理管线为：原始图片 → U²-Net 去背景 → 贴白底 → Resize(256) → CenterCrop(224) → Normalize → 模型推理 → Top-5 结果。', indent=True)

add_heading('3.3 主要工作和特色', 2)
add_para('本人在项目中完成的主要工作包括：', indent=True)
add_para('（1）数据集选型与分析：在多个公开数据集中（GTSRB 交通标志、Oxford-IIIT Pet 猫狗分类、Fruits-360 水果识别等）综合比较后，选择了 Fruits-360 作为本项目的基准数据集，原因是其类别数量适中（260 类）、数据规模大（18 万+张）、识别难度有层次（外观相似的不同品种水果构成挑战）。', indent=True)
add_para('（2）两阶段迁移学习训练策略的设计与实现：设计了"冻结训练 → 分层微调"的两阶段训练策略，并采用差异化学习率（fc > layer4 > layer3），使模型快速收敛并获得 98.80% 的高准确率。', indent=True)
add_para('（3）领域偏移问题的发现与解决：在真实场景测试中发现，模型在 Fruits-360 测试集上准确率达 98.80%，但面对网络真实照片时识别率骤降。通过系统分析，定位到问题根源是训练数据与实际应用之间的"领域偏移"（Domain Shift）——训练集均为纯白背景单水果，而真实图片背景复杂。通过引入 U²-Net 自动去背景模块，将识别管线改进为"抠图→贴白底→识别"，使得在真实网图上的 Top-1 正确率从 0/4 提升至 2/4，Top-3 命中率从 0/4 提升至 3/4。', indent=True)
add_para('（4）全面的模型评估体系：不仅计算了准确率，还生成了精确率/召回率/F1-score 分类报告、混淆矩阵可视化、易混淆类别分析等完整的评估诊断，能够精确回答"模型在哪里犯错"而非仅给出一个数字。', indent=True)
add_para('（5）Web 应用部署：基于 Flask 框架开发了完整的在线识别系统，集成了自动抠图、模型推理和 Top-5 结果展示，前后端交互流畅。', indent=True)
add_para('项目的主要特色在于：（a）从发现问题（领域偏移）到提出解决方案（抠图预处理）再到实验验证的完整分析闭环，体现了工程问题导向的思维方式；（b）两阶段训练策略保证了训练的稳定性和高效性；（c）对模型能力的清醒认知——准确率数字高不代表实际体验好，必须到真实场景中去检验。', indent=True)

add_heading('3.4 结果讨论', 2)

add_para('（1）训练结果分析', bold=True)
add_para('本项目最终在测试集上取得了 98.80% 的准确率，这是一个非常优异的结果。分析其原因：（a）ResNet-18 在 ImageNet 上预训练的权重为模型提供了强大的视觉特征提取基础；（b）两阶段训练策略避免了预训练权重被随机初始化的 FC 层"带坏"的问题；（c）数据增强（随机裁剪、翻转、旋转、颜色抖动）提升了模型的泛化能力。', indent=True)

add_para('值得注意的是，训练过程中第二阶段的准确率增长在后期趋于平缓（从 98.06% 到 98.80% 仅提升了 0.74 个百分点）。进一步分析分类报告发现，剩余的分类错误主要集中在以下三类问题上：', indent=True)
add_para('（a）数据集标签标注不一致：如 "BlackBerry 4" 与 "Blackberry 1/2/3" 实际为同一类水果但因大小写不一致被标注为不同类别，导致这些类别间的 F1-score 接近 0.00。这是数据质量而非模型能力的问题。', indent=True)
add_para('（b）类间视觉差异极小：不同编号的苹果（Apple Red 1/2/3）、不同品种的黑莓之间存在几乎难以用肉眼区分的细微差异，分类器在这些类别上的混淆是合理的。（c）数据集固有的性能天花板：Fruits-360 数据集的测试准确率在各论文中普遍集中在 97-99% 区间，本项目的 98.80% 已接近该数据集的最优水平。进一步更换模型架构（如 ResNet-50、EfficientNet）或引入注意力机制（如 CBAM），预期提升空间不超过 1 个百分点。', indent=True)

add_para('（2）领域偏移（Domain Shift）深入分析', bold=True)
add_para('本项目的一个核心发现是：高测试准确率不等于好的实际体验。模型在 Fruits-360 标准测试集上达到 98.80%，但面对从 Unsplash 下载的四张真实场景水果照片时，未加预处理的识别完全失败——所有四张图片均被错误分类（苹果树→Hazelnut、香蕉串→Pepper Yellow、橙子→Pepper Yellow、草莓→Blackberry）。', indent=True)
add_para('根本原因在于训练数据与实际应用之间的"领域偏移"（Domain Shift），具体表现为三个层面：', indent=True)
add_para('第一层——背景差异：Fruits-360 所有图片均为纯白背景，而真实图片有桌面、草地、天空、树枝等各种背景。模型在训练时未见过这些场景，容易将背景中的纹理混淆为分类特征。', indent=True)
add_para('第二层——构图差异：训练集中水果始终居中且占画面主导地位，而真实拍照中水果可能偏在角落或只占画面一小部分。', indent=True)
add_para('第三层——对象差异：训练集每张图只包含单个水果，而真实场景可能出现多水果（果盘、树上多颗果实），模型无法应对多目标场景。', indent=True)
add_para('为验证这些假设，本项目设计了三组对比实验。实验结果表明：引入 U²-Net 自动去背景可以有效解决第一层（背景差异）问题，抠图后在真实网图上的 Top-1 正确率从 0/4 提升至 2/4，Top-3 命中率从 0/4 提升至 3/4。但后两层问题（多目标、场景差异）无法通过单纯预处理解决，需要通过扩充带多样背景和多目标的训练数据，或采用目标检测（Object Detection）架构来解决。', indent=True)

add_para('（3）改进方向展望', bold=True)
add_para('基于以上分析，本项目的改进方向分为三个层次：', indent=True)
add_para('短期改进（预处理层面）：优化抠图模块的推理速度（当前 U²-Net 每次约 2-5 秒），尝试更轻量的分割模型（如 IS-Net）以减少延迟。', indent=True)
add_para('中期改进（数据层面）：在训练集中混入带有复杂背景的水果图片（如从网络抓取并人工标注），使模型直接学习区分水果与背景；对原始数据集进行标签清洗，合并语义相同但名称不一致的类别。', indent=True)
add_para('长期改进（架构层面）：将单标签分类扩展为多标签或目标检测任务，使模型能够同时识别图片中的多个水果并定位其位置，从而适应真实世界中的多目标场景。', indent=True)
add_para('综合来看，本项目的最重要认识是：在深度学习项目中，当数据质量已接近上限时，改进数据的收益远大于更换模型架构。"数据驱动"而非"模型驱动"的思维方式，是本次项目最宝贵的认知收获。', indent=True)

# =====================================================================
# 四、项目总结与收获
# =====================================================================
add_heading('四、项目总结与收获', 1)
add_heading('4.1 项目总结', 2)
add_para('本项目成功构建了一个基于 ResNet-18 迁移学习的水果智能识别系统。系统能够对 260 种水果和蔬菜进行自动分类，在 45,724 张测试集上达到 98.80% 的识别准确率。项目涵盖了从数据准备、数据增强、模型构建、两阶段训练、全面评估到 Web 部署的完整深度学习项目流程。', indent=True)
add_para('项目的创新点在于：（1）采用了两阶段训练策略（冻结训练→分层微调），在保证训练稳定性的同时充分挖掘预训练模型的潜力；（2）实现了完整的模型评估体系，通过混淆矩阵、易混淆类别分析等方法全面衡量模型性能；（3）将模型部署为 Web 应用，具备实际使用价值。', indent=True)
add_para('未来可以将本系统扩展到更多领域，如蔬菜品质分级、水果成熟度检测、病虫害识别等农业应用场景，也可以移植到移动端（Android/iOS）实现拍照即识别的便捷体验。', indent=True)

add_heading('4.2 心得体会', 2)
add_para('通过本次人工智能项目大作业的完整实践，我获得了以下方面的收获和体会：', indent=True)
add_para('第一，对深度学习基础理论有了更深入的理解。在项目实施前，我对卷积神经网络、迁移学习、梯度下降等概念只有模糊的认识。通过亲手编写训练循环代码（forward → loss → backward → step），我真正理解了神经网络的训练机制——前向传播计算预测值，损失函数衡量预测与真实的差距，反向传播计算梯度，优化器根据梯度更新参数。这四个步骤构成了所有深度学习训练的基石。', indent=True)
add_para('第二，理解了工程实践中的关键细节。数据增强不仅是"多加几张图"那么简单，随机裁剪模拟了不同的构图和视角，颜色抖动模拟了真实场景的光照变化，这些技巧的有效应用直接影响模型的泛化能力。分层学习率的设计源于对 ResNet 架构层次的理解：底层检测边缘纹理（通用特征），高层检测形状组合（任务相关特征）。这种"底层冻结、高层微调"的策略是迁移学习的核心思想。', indent=True)
add_para('第三，体会到了实验迭代的重要性。在项目过程中遇到了多个实际问题：Windows 系统下 DataLoader 多进程兼容性问题（需将 num_workers 设为 0）、GPU 显存限制对 Batch Size 选择的影响、模型训练过程中的输出缓冲导致看不到实时进度等。这些问题都是在理论学习中不会遇到的，只有在动手实践中才能真正积累解决经验。', indent=True)
add_para('第四，深刻体会到了"数据驱动"的思维方式。在项目初期，模型在 Fruits-360 测试集上获得 98.80% 的高准确率时，我曾一度认为项目已经完成。但当我将模型部署为 Web 应用并用网上搜到的真实水果照片测试时，识别效果却远远不如预期——几乎所有带背景的图片都被错误分类。这个反差让我深刻认识到：测试集的高分不代表真实世界中的好体验。通过系统排查，我定位到问题根源是数据集领域的 "领域偏移"（Domain Shift）：Fruits-360 的图片全是在摄影棚内拍摄的白底单果图，而现实世界中的水果照片背景复杂、构图多样。基于此认知，我引入了 U²-Net 自动抠图模块作为预处理，有效缓解了背景差异带来的识别下降。更重要的是，我认识到在当前数据集上，剩余的性能瓶颈不在模型架构而在于数据质量和分布——"高质量数据往往比复杂模型更关键"。这个来自实践的认识远比课本上的公式更加深刻，对我未来从事 AI 研究具有重要的指导意义。', indent=True)
add_para('第五，体会到了分析和表达的重要性。一个好的 AI 项目不仅需要技术实现，还需要能够清晰地表达自己的发现和思考。项目中通过对比实验（有/无抠图）量化了领域偏移的影响程度，通过混淆矩阵定位了模型的薄弱类别，通过分类报告分析了错误的根本原因——这些分析让答辩汇报"有据可依"，而不仅仅是堆砌技术名词。', indent=True)
add_para('本次项目为将来在人工智能方向的深入学习和研究打下了坚实的基础。掌握的知识和技能不仅限于水果识别这一个任务，迁移学习的方法论可以应用到任何图像分类场景，为研究生阶段的科研工作做好了准备。', indent=True)

# ===== 保存 =====
output_path = os.path.join(BASE, '水果智能识别系统_设计报告_v2.docx')
doc.save(output_path)
print(f'报告已保存: {output_path}')
